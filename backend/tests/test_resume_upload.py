import io
import unittest
from unittest.mock import AsyncMock, MagicMock, patch

import docx
from fastapi.testclient import TestClient

from main import app


class TestResumeUpload(unittest.TestCase):
    def setUp(self):
        self.client = TestClient(app)

    @patch("services.resume_service.upload_file_to_supabase", new_callable=AsyncMock)
    @patch("services.resume_service.pdfplumber.open")
    def test_upload_valid_pdf(self, mock_pdf_open, mock_supabase):
        mock_supabase.side_effect = lambda user_id, file_ref_id, filename, file_bytes, content_type: f"{user_id}/{file_ref_id}_{filename}"

        mock_page = MagicMock()
        mock_page.extract_text.return_value = "John Doe\nSoftware Engineer\nPython FastAPI"
        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_pdf.__enter__.return_value = mock_pdf
        mock_pdf_open.return_value = mock_pdf

        pdf_bytes = b"%PDF-1.4 fake pdf content for testing"
        files = {"file": ("resume.pdf", pdf_bytes, "application/pdf")}
        data = {"user_id": "user_123"}

        response = self.client.post("/resume/upload", files=files, data=data)
        self.assertEqual(response.status_code, 200)

        json_resp = response.json()
        self.assertIn("file_reference_id", json_resp)
        self.assertIn("text", json_resp)
        self.assertEqual(json_resp["filename"], "resume.pdf")
        self.assertTrue(json_resp["storage_path"].startswith("user_123/"))
        self.assertTrue(json_resp["storage_path"].endswith("_resume.pdf"))
        self.assertIn("John Doe", json_resp["text"])

    @patch("services.resume_service.upload_file_to_supabase", new_callable=AsyncMock)
    def test_upload_valid_docx(self, mock_supabase):
        mock_supabase.side_effect = lambda user_id, file_ref_id, filename, file_bytes, content_type: f"{user_id}/{file_ref_id}_{filename}"

        doc = docx.Document()
        doc.add_paragraph("Jane Smith")
        doc.add_paragraph("Senior Python Developer")
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_bytes = docx_io.getvalue()

        files = {
            "file": (
                "resume.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        data = {"user_id": "user_123"}

        response = self.client.post("/resume/upload", files=files, data=data)
        self.assertEqual(response.status_code, 200)

        json_resp = response.json()
        self.assertIn("file_reference_id", json_resp)
        self.assertTrue(json_resp["storage_path"].startswith("user_123/"))
        self.assertTrue(json_resp["storage_path"].endswith("_resume.docx"))
        self.assertIn("Jane Smith", json_resp["text"])
        self.assertIn("Senior Python Developer", json_resp["text"])

    def test_upload_corrupt_pdf_returns_422(self):
        corrupt_bytes = b"NOT_A_REAL_PDF_DATA_HEADER"
        files = {"file": ("corrupt_resume.pdf", corrupt_bytes, "application/pdf")}

        response = self.client.post("/resume/upload", files=files)
        self.assertEqual(response.status_code, 422)
        json_resp = response.json()
        self.assertIn("detail", json_resp)
        self.assertIn("Corrupt", json_resp["detail"])

    def test_upload_corrupt_docx_returns_422(self):
        corrupt_bytes = b"NOT_A_REAL_DOCX_FILE"
        files = {
            "file": (
                "corrupt_resume.docx",
                corrupt_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }

        response = self.client.post("/resume/upload", files=files)
        self.assertEqual(response.status_code, 422)
        json_resp = response.json()
        self.assertIn("detail", json_resp)
        self.assertIn("Corrupt", json_resp["detail"])

    def test_upload_unsupported_file_returns_422(self):
        txt_bytes = b"Some random text file"
        files = {"file": ("resume.png", txt_bytes, "image/png")}

        response = self.client.post("/resume/upload", files=files)
        self.assertEqual(response.status_code, 422)
        json_resp = response.json()
        self.assertIn("Unsupported file format", json_resp["detail"])


if __name__ == "__main__":
    unittest.main()
